import docx

doc = docx.Document('Week05_BaoCaoNhom.docx')

members = [
  "Nguyễn Trần Gia Bảo",
  "Lê Văn Minh Trí",
  "Nguyễn Duy Quang",
  "Huỳnh Minh Huân",
  "Đỗ Minh Gia Bảo",
  "(Không áp dụng)"
]

# Helper function to find a table by a keyword in its first row
def find_table(keyword):
    for idx, t in enumerate(doc.tables):
        try:
            if keyword in t.cell(0, 0).text or keyword in t.cell(0, 1).text or keyword in t.cell(0, 2).text:
                return t
        except:
            pass
    return None

# Table 0: Info
t_info = doc.tables[0]
if t_info:
    t_info.cell(0, 1).text = "Auralis (Demo)"
    t_info.cell(1, 1).text = "FER202"
    t_info.cell(4, 1).text = "Tuần 05 — React Bootstrap & Responsive"

# Table 1: Phân công (STT | Họ và tên | Nhiệm vụ...)
t_phan_cong = find_table("Nhiệm vụ được giao")
if t_phan_cong:
    tasks_phan_cong = [
      "Header → Navbar + Nav + Container. Rebuild SearchBar dùng Form, InputGroup.",
      "HomePage Banner → Jumbotron style; CategoryList → Row/Col",
      "TrackCard → Card + Badge + Button",
      "TrackGrid → Row + Col với breakpoints xs, sm, md, lg, xl",
      "Test responsive toàn bộ app tại 375px, 768px, 1280px"
    ]
    for i in range(5):
        t_phan_cong.cell(i+1, 1).text = members[i]
        t_phan_cong.cell(i+1, 2).text = tasks_phan_cong[i]
        t_phan_cong.cell(i+1, 3).text = "Đạt"
        t_phan_cong.cell(i+1, 4).text = "100%"

# Table 2: Chủ trì (Người chủ trì tuần này)
t_chu_tri = find_table("Người chủ trì")
if t_chu_tri:
    t_chu_tri.cell(0, 1).text = members[0]
    t_chu_tri.cell(1, 1).text = "https://github.com/..."

# Table 3: Danh sách công việc (Công việc cần thực hiện | T.thái)
t_cong_viec = find_table("Công việc cần thực hiện")
if t_cong_viec:
    t_cong_viec.cell(1, 2).text = "[x]"
    t_cong_viec.cell(1, 3).text = members[0]
    
    t_cong_viec.cell(2, 2).text = "[x]"
    t_cong_viec.cell(2, 3).text = members[3]
    
    t_cong_viec.cell(3, 2).text = "[x]"
    t_cong_viec.cell(3, 3).text = members[2]
    
    t_cong_viec.cell(4, 2).text = "[x]"
    t_cong_viec.cell(4, 3).text = members[1]
    
    t_cong_viec.cell(5, 2).text = "[x]"
    t_cong_viec.cell(5, 3).text = members[2]
    
    t_cong_viec.cell(6, 2).text = "[x]"
    t_cong_viec.cell(6, 3).text = members[4]

# Paragraphs: mô tả
for p in doc.paragraphs:
    if "mô tả approach đã chọn" in p.text:
        p.text = "Trong tuần này, nhóm đã thực hiện đại tu kiến trúc UI toàn diện: chuyển từ các utility classes của TailwindCSS sang hệ thống Component chuẩn mực của React-Bootstrap. Xóa bỏ các div tĩnh chứa class Flexbox, thay bằng Container, Row và Col để responsive tự động. Các thẻ tĩnh được thay bằng Card, Badge, Form. Khó khăn lớn nhất là màu mặc định của Bootstrap bị sáng, không hợp với Dark Theme. Giải pháp là ghi đè các biến CSS của Bootstrap (--bs-dark, --bs-primary) trong index.css để giữ được chất màu tối."

# Table 5: Yêu cầu đầu ra
t_out = find_table("Yêu cầu đầu ra")
if t_out:
    t_out.cell(1, 2).text = "[x]"
    t_out.cell(1, 3).text = "src/components/, src/App.tsx"
    t_out.cell(1, 4).text = "Không còn class TailwindCSS layout"

    t_out.cell(2, 2).text = "[x]"
    t_out.cell(2, 3).text = "docs/week05-responsive/"
    t_out.cell(2, 4).text = "Các Screenshot đã chụp trên 3 breakpoints"

    t_out.cell(3, 2).text = "[x]"
    t_out.cell(3, 3).text = "Week05_BaoCaoNhom.docx"
    t_out.cell(3, 4).text = "Sử dụng Navbar, Container, Row, Col, Card, Badge, Form, Button"
    
    t_out.cell(4, 1).text = "3/3"

# Table 6: Lỗi
t_err = find_table("Mô tả lỗi")
if t_err:
    t_err.cell(1, 1).text = "Menu Navbar không thể tự động đổ xuống trên màn hình nhỏ (Mobile)"
    t_err.cell(1, 2).text = "Code thiếu thẻ Navbar.Toggle và Navbar.Collapse"
    t_err.cell(1, 3).text = "Thêm các thẻ tương ứng bọc lấy Nav"
    t_err.cell(1, 4).text = "Đã fix"

    t_err.cell(2, 1).text = "Màu thẻ Card bị trắng toát không phù hợp với Dark Theme"
    t_err.cell(2, 2).text = "Thuộc tính mặc định của Bootstrap Card là light background"
    t_err.cell(2, 3).text = "Set thuộc tính bg='dark' và dùng --bs-dark custom"
    t_err.cell(2, 4).text = "Đã fix"

doc.save("Week05_BaoCaoNhom_V4.docx")
